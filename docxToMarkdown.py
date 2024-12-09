from docx import Document


def docx_to_markdown(file_name, output_file_name):
    """
    Converts a Word document (.docx) into Markdown-style format and saves it as a new Word document.
    This makes the report template more understandable for the LLM.

    Parameters:
    file_name (str): Input Word document path.
    output_file_name (str): Output Word document path.

    Converts headings, bullet points, numbered lists, horizontal lines, and paragraphs to Markdown syntax.
    Saves the converted content to a new Word document and prints a confirmation message.
    """


    # Load the Word document
    doc = Document(file_name)
    
    # Create a new Word document for saving
    new_doc = Document()


    markdown_text = ""

    for paragraph in doc.paragraphs:
        # Check the style of each paragraph and convert to Markdown
        if paragraph.style.name.startswith("Heading 1"):
            markdown_text += f"# {paragraph.text}\n\n"
        elif paragraph.style.name.startswith("Heading 2"):
            markdown_text += f"## {paragraph.text}\n\n"
        elif paragraph.style.name.startswith("Heading 3"):
            markdown_text += f"### {paragraph.text}\n\n"
        elif paragraph.style.name.startswith("Heading 4"):
            markdown_text += f"#### {paragraph.text}\n\n"
        elif paragraph.style.name.startswith("List Bullet"):
            markdown_text += f"- {paragraph.text}\n"
        elif paragraph.style.name.startswith("List Number"):
            markdown_text += f"1. {paragraph.text}\n"
        elif "---" in paragraph.text or "Horizontal Line" in paragraph.style.name:
            # Horizontal line
            markdown_text += "---\n\n"
        else:
            # Regular paragraph
            markdown_text += f"{paragraph.text}\n\n"
    
    new_doc.add_paragraph(markdown_text)
    # Save the new Word document
    new_doc.save(output_file_name)
    print(f"Document saved as {output_file_name}")