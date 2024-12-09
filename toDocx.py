import markdown
from docx import Document
from bs4 import BeautifulSoup

def markdown_to_docx(md_text, file_name):

    """
    Converts Markdown text to a Word document (.docx). 
    This can convert the LLM report output into a downloadable .docx file.

    Parameters:
    md_text (str): Markdown content to convert.
    file_name (str): Path to save the resulting Word document.

    Converts headings, paragraphs, bold/italic text, and lists into corresponding Word styles. Saves the document at the specified file path.

    Example:
    markdown_to_docx("# Heading 1\nText", "output.docx")
    """

    # Convert Markdown to HTML
    html = markdown.markdown(md_text)
    
    # Use BeautifulSoup to parse the HTML
    soup = BeautifulSoup(html, 'html.parser')
    
    # Create a new Word document
    doc = Document()
    
    # Helper function to add paragraphs with formatting
    def add_paragraph_with_format(text, bold=False, italic=False):
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(text)
        run.bold = bold
        run.italic = italic

    # Iterate over the elements in the HTML
    for element in soup.contents:
        if element.name == 'h1':
            doc.add_heading(element.text, level=1)
        elif element.name == 'h2':
            doc.add_heading(element.text, level=2)
        elif element.name == 'h3':
            doc.add_heading(element.text, level=3)
        elif element.name == 'p':
            add_paragraph_with_format(element.text)
        elif element.name == 'strong':
            add_paragraph_with_format(element.text, bold=True)
        elif element.name == 'em':
            add_paragraph_with_format(element.text, italic=True)
        elif element.name == 'ul':
            for li in element.find_all('li'):
                doc.add_paragraph(li.text, style='List Bullet')
        elif element.name == 'ol':
            for li in element.find_all('li'):
                doc.add_paragraph(li.text, style='List Number')
    
    # Save the Word document
    doc.save(file_name)
    print(f"Document saved as {file_name}")
